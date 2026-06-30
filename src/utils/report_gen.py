import io
import logging
from datetime import datetime
from typing import Dict, Any, List

logger = logging.getLogger("ReportGen")

class ReportGenerator:
    def __init__(self):
        pass

    def generate_pdf_report(self, 
                            kpis: Dict[str, Any], 
                            executive_summary: str, 
                            recommendations: List[Dict[str, str]], 
                            model_perf: Dict[str, Any]) -> bytes:
        """
        Generates a PDF report using fpdf2 and returns the bytes.
        """
        try:
            from fpdf import FPDF
            
            class PDF(FPDF):
                def header(self):
                    self.set_font('Helvetica', 'B', 15)
                    self.set_text_color(44, 62, 80)
                    self.cell(0, 10, 'SentiVerse AI - Executive Intelligence Report', 0, 1, 'C')
                    self.ln(5)
                    
                def footer(self):
                    self.set_y(-15)
                    self.set_font('Helvetica', 'I', 8)
                    self.set_text_color(127, 140, 141)
                    self.cell(0, 10, f'Page {self.page_no()} | Generated on {datetime.now().strftime("%Y-%m-%d %H:%M")}', 0, 0, 'C')

            pdf = PDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            
            # Subtitle
            pdf.set_font('Helvetica', 'B', 12)
            pdf.set_text_color(52, 73, 94)
            pdf.cell(0, 8, f'Dataset: {kpis.get("dataset_name", "Social Media Feed")}', 0, 1, 'L')
            pdf.cell(0, 8, f'Total Posts Analyzed: {kpis.get("total_posts", 0)}', 0, 1, 'L')
            pdf.ln(5)
            
            # Draw KPI Boxes (Reputation & CSAT)
            pdf.set_fill_color(240, 243, 244)
            pdf.rect(10, pdf.get_y(), 90, 25, 'F')
            pdf.rect(110, pdf.get_y(), 90, 25, 'F')
            
            pdf.set_font('Helvetica', 'B', 10)
            pdf.set_text_color(127, 140, 141)
            pdf.set_xy(15, pdf.get_y() + 2)
            pdf.cell(80, 5, 'BRAND REPUTATION SCORE', 0, 0, 'L')
            pdf.set_xy(115, pdf.get_y())
            pdf.cell(80, 5, 'ESTIMATED CSAT SCORE', 0, 1, 'L')
            
            pdf.set_font('Helvetica', 'B', 16)
            pdf.set_text_color(46, 204, 113) # Green
            pdf.set_xy(15, pdf.get_y() + 2)
            pdf.cell(80, 10, f'{kpis.get("brand_reputation", 50.0)} / 100', 0, 0, 'L')
            pdf.set_xy(115, pdf.get_y())
            pdf.cell(80, 10, f'{kpis.get("csat", 50.0)} %', 0, 1, 'L')
            
            pdf.set_xy(10, pdf.get_y() + 10)
            
            # Executive Summary
            pdf.set_font('Helvetica', 'B', 12)
            pdf.set_text_color(44, 62, 80)
            pdf.cell(0, 10, 'Executive Summary', 0, 1, 'L')
            pdf.set_font('Helvetica', '', 10)
            pdf.set_text_color(51, 51, 51)
            pdf.multi_cell(0, 6, executive_summary)
            pdf.ln(5)
            
            # Recommendations
            pdf.set_font('Helvetica', 'B', 12)
            pdf.set_text_color(44, 62, 80)
            pdf.cell(0, 10, 'Product & Marketing Recommendations', 0, 1, 'L')
            
            for idx, rec in enumerate(recommendations):
                pdf.set_font('Helvetica', 'B', 10)
                pdf.set_text_color(230, 126, 34) # Orange
                pdf.cell(0, 6, f'{idx+1}. {rec.get("category", "Action Item")}', 0, 1, 'L')
                
                pdf.set_font('Helvetica', 'I', 9)
                pdf.set_text_color(127, 140, 141)
                pdf.cell(0, 5, f'Trigger: {rec.get("issue", "")}', 0, 1, 'L')
                
                pdf.set_font('Helvetica', '', 10)
                pdf.set_text_color(51, 51, 51)
                pdf.multi_cell(0, 5, f'Recommendation: {rec.get("action", "")}')
                pdf.ln(3)
                
            # Model Performance
            pdf.set_font('Helvetica', 'B', 12)
            pdf.set_text_color(44, 62, 80)
            pdf.cell(0, 10, 'Model Classification Performance Summary', 0, 1, 'L')
            
            pdf.set_font('Helvetica', '', 10)
            for model_name, acc in model_perf.items():
                pdf.cell(80, 6, f'- {model_name}:', 0, 0, 'L')
                pdf.cell(0, 6, f'Accuracy: {acc:.4f}' if isinstance(acc, float) else f'{acc}', 0, 1, 'L')

            # Output bytes
            buffer = io.BytesIO()
            # fpdf2 output(dest='S') returns bytes
            pdf_bytes = pdf.output()
            if isinstance(pdf_bytes, bytearray):
                return bytes(pdf_bytes)
            return pdf_bytes
            
        except Exception as e:
            logger.error(f"Failed to generate PDF report: {e}. Falling back to plain text PDF.")
            # Simple text output in bytes
            return self.generate_text_report(kpis, executive_summary, recommendations, model_perf).encode('utf-8')

    def generate_word_report(self, 
                             kpis: Dict[str, Any], 
                             executive_summary: str, 
                             recommendations: List[Dict[str, str]], 
                             model_perf: Dict[str, Any]) -> bytes:
        """
        Generates a Word report using python-docx and returns the bytes.
        """
        try:
            from docx import Document
            
            doc = Document()
            doc.add_heading('SentiVerse AI - Executive Intelligence Report', 0)
            
            doc.add_paragraph(f'Dataset: {kpis.get("dataset_name", "Social Media Feed")}')
            doc.add_paragraph(f'Total Posts Analyzed: {kpis.get("total_posts", 0)}')
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            
            doc.add_heading('Key Performance Indicators', level=1)
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Value'
            
            row_cells_1 = table.add_row().cells
            row_cells_1[0].text = 'Brand Reputation Score'
            row_cells_1[1].text = f'{kpis.get("brand_reputation", 50.0)} / 100'
            
            row_cells_2 = table.add_row().cells
            row_cells_2[0].text = 'Estimated CSAT Score'
            row_cells_2[1].text = f'{kpis.get("csat", 50.0)}%'
            
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(executive_summary)
            
            doc.add_heading('Actionable Product Suggestions', level=1)
            for idx, rec in enumerate(recommendations):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f'{rec.get("category", "Action")}: ').bold = True
                p.add_run(f'{rec.get("action", "")} (Based on: {rec.get("issue", "")})')
                
            doc.add_heading('Sentiment Prediction Models Benchmarks', level=1)
            for model_name, acc in model_perf.items():
                doc.add_paragraph(f'{model_name}: Accuracy = {acc:.4f}' if isinstance(acc, float) else f'{model_name}: {acc}')
                
            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Failed to generate Word report: {e}. Falling back to plain text.")
            return self.generate_text_report(kpis, executive_summary, recommendations, model_perf).encode('utf-8')

    def generate_text_report(self, 
                             kpis: Dict[str, Any], 
                             executive_summary: str, 
                             recommendations: List[Dict[str, str]], 
                             model_perf: Dict[str, Any]) -> str:
        """
        Fallback plain text report.
        """
        report = (
            f"SENTIVERSE AI EXECUTIVE REPORT\n"
            f"=============================\n"
            f"Dataset Name: {kpis.get('dataset_name', 'Social Media Feed')}\n"
            f"Total Posts: {kpis.get('total_posts', 0)}\n"
            f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"
            f"KPIs:\n"
            f"- Brand Reputation Score: {kpis.get('brand_reputation', 50.0)} / 100\n"
            f"- CSAT Score: {kpis.get('csat', 50.0)}%\n\n"
            f"Executive Summary:\n"
            f"{executive_summary}\n\n"
            f"Recommendations:\n"
        )
        for idx, rec in enumerate(recommendations):
            report += f"{idx+1}. [{rec.get('category')}] {rec.get('action')} (Trigger: {rec.get('issue')})\n"
            
        report += "\nModel Performance Summary:\n"
        for name, acc in model_perf.items():
            report += f"- {name}: {acc:.4f}\n" if isinstance(acc, float) else f"- {name}: {acc}\n"
            
        return report
